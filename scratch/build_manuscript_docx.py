"""Build the final competition DOCX from the Markdown manuscript, following the
Prêmio Jovem Cientista edital: A4, Arial 12, line spacing 1.5, Portuguese,
figures embedded, tables formatted.

Run: python scratch/build_manuscript_docx.py
"""
from __future__ import annotations
import os, re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

SRC = "docs/manuscrito/PrimeVarClass_artigo_honesto.md"
MDDIR = os.path.dirname(SRC)
OUT = "docs/manuscrito/PrimeVarClass_Artigo_Premio_Jovem_Cientista.docx"
FONT = "Arial"

doc = Document()

# ---- base style: Arial 12, spacing 1.5 ----
normal = doc.styles["Normal"]
normal.font.name = FONT
normal.font.size = Pt(12)
normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
pf = normal.paragraph_format
pf.line_spacing = 1.5
pf.space_after = Pt(4)

for hs, sz in [("Heading 1", 14), ("Heading 2", 12), ("Heading 3", 12), ("Title", 16)]:
    try:
        st = doc.styles[hs]
        st.font.name = FONT
        st.font.size = Pt(sz)
        st.font.color.rgb = RGBColor(0, 0, 0)
    except KeyError:
        pass

# ---- A4 + margins ----
sec = doc.sections[0]
sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)
for m in ("top_margin", "bottom_margin"):
    setattr(sec, m, Cm(2.5))
for m in ("left_margin", "right_margin"):
    setattr(sec, m, Cm(3.0))
CONTENT_CM = 13.0

INLINE = re.compile(r"(\*\*.+?\*\*|\*[^*]+?\*|`[^`]+?`)")


def add_runs(p, text):
    for part in INLINE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = p.add_run(part[2:-2]); r.bold = True
        elif part.startswith("*") and part.endswith("*"):
            r = p.add_run(part[1:-1]); r.italic = True
        elif part.startswith("`") and part.endswith("`"):
            r = p.add_run(part[1:-1]); r.font.name = "Consolas"
        else:
            r = p.add_run(part)
        r.font.name = FONT if not (part.startswith("`")) else "Consolas"


def col_aligns(sep):
    al = []
    for c in [x.strip() for x in sep.strip().strip("|").split("|")]:
        if c.endswith(":") and c.startswith(":"):
            al.append(WD_ALIGN_PARAGRAPH.CENTER)
        elif c.endswith(":"):
            al.append(WD_ALIGN_PARAGRAPH.RIGHT)
        else:
            al.append(WD_ALIGN_PARAGRAPH.LEFT)
    return al


def add_table(rows, aligns):
    cells0 = [c.strip() for c in rows[0].strip().strip("|").split("|")]
    n = len(cells0)
    t = doc.add_table(rows=len(rows), cols=n)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ri, row in enumerate(rows):
        cells = [c.strip() for c in row.strip().strip("|").split("|")]
        for ci in range(n):
            cell = t.cell(ri, ci)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_after = Pt(2)
            p.alignment = aligns[ci] if ci < len(aligns) else WD_ALIGN_PARAGRAPH.LEFT
            txt = cells[ci] if ci < len(cells) else ""
            add_runs(p, txt)
            for run in p.runs:
                run.font.size = Pt(10)
                if ri == 0:
                    run.bold = True
    doc.add_paragraph()


lines = open(SRC, encoding="utf-8").read().split("\n")
i = 0
while i < len(lines):
    ln = lines[i]
    s = ln.strip()
    if not s:
        i += 1; continue
    # horizontal rule
    if s == "---":
        i += 1; continue
    # images
    m = re.match(r"^!\[.*?\]\((.+?)\)\s*$", s)
    if m:
        path = os.path.normpath(os.path.join(MDDIR, m.group(1)))
        if os.path.exists(path):
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(path, width=Cm(CONTENT_CM))
        i += 1; continue
    # captions (Figura/Tabela) — italic centered small
    if re.match(r"^\*\*(Figura|Tabela)\b", s):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_runs(p, s)
        for r in p.runs:
            r.font.size = Pt(10); r.italic = True
        i += 1; continue
    # headings
    if s.startswith("# "):
        p = doc.add_paragraph(s[2:].strip(), style="Title"); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        i += 1; continue
    if s.startswith("## "):
        doc.add_paragraph(s[3:].strip(), style="Heading 1"); i += 1; continue
    if s.startswith("### "):
        doc.add_paragraph(s[4:].strip(), style="Heading 2"); i += 1; continue
    # blockquote front-matter
    if s.startswith(">"):
        block = []
        while i < len(lines) and lines[i].strip().startswith(">"):
            block.append(lines[i].strip()[1:].strip()); i += 1
        for b in block:
            if not b:
                continue
            p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
            add_runs(p, b)
            for r in p.runs:
                r.font.size = Pt(11)
        doc.add_paragraph()
        continue
    # tables
    if s.startswith("|") and i + 1 < len(lines) and re.match(r"^\s*\|?[\s:-]+\|", lines[i + 1]):
        tbl = [lines[i]]
        aligns = col_aligns(lines[i + 1])
        i += 2
        while i < len(lines) and lines[i].strip().startswith("|"):
            tbl.append(lines[i]); i += 1
        add_table(tbl, aligns)
        continue
    # bullet list
    if s.startswith("- "):
        p = doc.add_paragraph(style="List Bullet")
        add_runs(p, s[2:].strip())
        i += 1; continue
    # normal paragraph (justified)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_runs(p, s)
    i += 1

doc.save(OUT)
print("saved:", OUT)
print("paragraphs:", len(doc.paragraphs), "tables:", len(doc.tables))
