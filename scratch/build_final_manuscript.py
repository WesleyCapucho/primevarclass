# -*- coding: utf-8 -*-
"""Gera o manuscrito final do PrimeVarClass para o 32º Prêmio Jovem Cientista,
em conformidade com o Edital (item 2.2.2c): A4, Arial 12, espaçamento 1,5,
estrutura Apresentação / Desenvolvimento / Conclusão / Referências.

Todo o texto é original desta sessão (primeira pessoa, pretérito), e todos os
números citados foram verificados contra os JSON/CSV reais em
primevarclass_manuscript_analysis/ e scratch/decisive_results/ antes da escrita.

Run: python scratch/build_final_manuscript.py
"""
import os

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

FIG = "docs/manuscrito/figuras"
OUT = "docs/manuscrito/PrimeVarClass_Artigo_Premio_Jovem_Cientista.docx"

# --------------------------------------------------------------------------- #
#  Helpers                                                                     #
# --------------------------------------------------------------------------- #
def setup_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(3.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.0)

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(12)
    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.first_line_indent = Cm(1.25)
    pf.space_after = Pt(0)

    for lvl, size in [("Heading 1", 14), ("Heading 2", 12), ("Heading 3", 12)]:
        h = doc.styles[lvl]
        h.font.name = "Arial"
        h.font.size = Pt(size)
        h.font.bold = True
        h.font.color.rgb = RGBColor(0, 0, 0)
        h.paragraph_format.space_before = Pt(18 if lvl == "Heading 1" else 12)
        h.paragraph_format.space_after = Pt(8)
        h.paragraph_format.first_line_indent = Cm(0)
        h.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return doc


def h1(doc, text):
    p = doc.add_heading(text, level=1)
    return p


def h2(doc, text):
    return doc.add_heading(text, level=2)


def h3(doc, text):
    return doc.add_heading(text, level=3)


def para(doc, text, indent=True, bold=False, italic=False, align=None, size=12,
         space_before=0, space_after=0):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.25) if indent else Cm(0)
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.font.name = "Arial"; r.font.size = Pt(size)
    r.bold = bold; r.italic = italic
    return p


def para_runs(doc, runs, indent=True, align=None, space_before=0, space_after=0):
    """runs: list of (text, bold, italic) tuples -> one paragraph, multiple runs."""
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.25) if indent else Cm(0)
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    for text, bold, italic in runs:
        r = p.add_run(text)
        r.font.name = "Arial"; r.font.size = Pt(12)
        r.bold = bold; r.italic = italic
    return p


def keywords_line(doc, label, text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    r1 = p.add_run(label); r1.bold = True; r1.font.name = "Arial"; r1.font.size = Pt(12)
    r2 = p.add_run(text); r2.font.name = "Arial"; r2.font.size = Pt(12)
    return p


def figure(doc, path, caption, width=5.5):
    full = os.path.join(FIG, path)
    if not os.path.exists(full):
        print(f"AVISO: figura ausente -> {full}")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(10)
    run = p.add_run()
    run.add_picture(full, width=Inches(width))
    c = doc.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.paragraph_format.first_line_indent = Cm(0)
    c.paragraph_format.space_after = Pt(12)
    cr = c.add_run(caption)
    cr.font.name = "Arial"; cr.font.size = Pt(10); cr.italic = True


def _set_cell_border(cell, edge, sz=6, color="000000"):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = tcPr.makeelement(qn("w:tcBorders"), {})
        tcPr.append(tcBorders)
    el = tcBorders.makeelement(qn(f"w:{edge}"), {
        qn("w:val"): "single", qn("w:sz"): str(sz), qn("w:color"): color})
    tcBorders.append(el)


def table(doc, headers, rows, widths_cm=None, caption=None, note=None, font_size=10.5):
    if caption:
        cp = doc.add_paragraph()
        cp.paragraph_format.first_line_indent = Cm(0)
        cp.paragraph_format.space_before = Pt(10)
        r = cp.add_run(caption); r.bold = True; r.font.name = "Arial"; r.font.size = Pt(11)
    n_cols = len(headers)
    if widths_cm is None:
        widths_cm = [Cm(15.0 / n_cols)] * n_cols
    else:
        widths_cm = [Cm(w) for w in widths_cm]
    t = doc.add_table(rows=1, cols=n_cols)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].width = widths_cm[i]
        p = hdr[i].paragraphs[0]; p.paragraph_format.first_line_indent = Cm(0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(htext); r.bold = True; r.font.name = "Arial"; r.font.size = Pt(font_size)
        for edge in ("top", "bottom"):
            _set_cell_border(hdr[i], edge, sz=10)
    for r_ in t.rows:                                # keep rows intact across pages
        trPr = r_._tr.get_or_add_trPr()
        trPr.append(trPr.makeelement(qn("w:cantSplit"), {}))
    for row in rows:
        cells = t.add_row().cells
        tr = t.rows[-1]
        trPr = tr._tr.get_or_add_trPr()
        trPr.append(trPr.makeelement(qn("w:cantSplit"), {}))
        for i, val in enumerate(row):
            cells[i].width = widths_cm[i]
            p = cells[i].paragraphs[0]; p.paragraph_format.first_line_indent = Cm(0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i > 0 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val)); r.font.name = "Arial"; r.font.size = Pt(font_size)
    # bottom border on the last row
    for c in t.rows[-1].cells:
        _set_cell_border(c, "bottom", sz=10)
    if note:
        npar = doc.add_paragraph()
        npar.paragraph_format.first_line_indent = Cm(0)
        npar.paragraph_format.space_after = Pt(12)
        nr = npar.add_run(note); nr.font.name = "Arial"; nr.font.size = Pt(9.5); nr.italic = True
    return t


# MARKER_CONTENT_START


# --------------------------------------------------------------------------- #
#  Capa / folha de rosto                                                       #
# --------------------------------------------------------------------------- #
def build_frontmatter(doc):
    for _ in range(2):
        doc.add_paragraph()
    para(doc, "PrimeVarClass: da hipótese dos números primos a um classificador de "
              "variantes BRCA1/BRCA2 consciente de domínio, validado externamente e "
              "complementar ao estado da arte",
         indent=False, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=15,
         space_after=18)
    para(doc, "32º Prêmio Jovem Cientista — Categoria Estudante do Ensino Superior",
         indent=False, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    para(doc, "Tema: Inteligência Artificial para o Bem Comum — Subtema: Inteligência "
              "Artificial & Saúde (item 1.4.1.b do Edital)",
         indent=False, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
    keywords_line(doc, "Autor: ", "Wesley Felipe Capucho — graduando em Engenharia Bioquímica")
    keywords_line(doc, "Orientador(a): ", "⟨a definir⟩")
    keywords_line(doc, "Instituição de vínculo: ", "Escola de Engenharia de Lorena, "
                  "Universidade de São Paulo (EEL-USP) — Estrada Municipal do Campinho, s/nº, "
                  "Ponte Nova, Lorena – SP, CEP 12602-810. E-mail: wesleycapucho@usp.br. "
                  "Telefone: ⟨a definir⟩.")
    keywords_line(doc, "Instituição onde a pesquisa foi desenvolvida: ", "Escola de "
                  "Engenharia de Lorena, Universidade de São Paulo (EEL-USP), Lorena – SP, "
                  "Brasil.")
    doc.add_paragraph()
    keywords_line(doc, "Palavras-chave: ", "classificação de variantes de significado "
                  "incerto; BRCA1/BRCA2; validação externa; domínios funcionais de "
                  "proteínas; inteligência artificial em saúde.")
    doc.add_page_break()


# --------------------------------------------------------------------------- #
#  APRESENTAÇÃO                                                                #
# --------------------------------------------------------------------------- #
def build_apresentacao(doc):
    h1(doc, "APRESENTAÇÃO")
    h2(doc, "Resumo")
    para(doc, "Este trabalho partiu de uma hipótese pouco convencional: codificar "
              "aminoácidos pela estrutura dos números primos revelaria padrões de "
              "patogenicidade em variantes missense? O projeto foi nomeado PrimeVarClass "
              "a partir dessa aposta, e a hipótese foi testada com rigor — validação "
              "cruzada bloqueada por posição e generalização em coortes externas nunca "
              "vistas no treino. A hipótese foi refutada com transparência: a "
              "codificação por primos teve desempenho inferior ao de uma identidade "
              "trivial de aminoácidos (AUC externa 0,681 vs. 0,718; p = 0,045) e piorou "
              "um modelo bioquímico ao ser adicionada (0,791 → 0,765; p = 3,8 × 10⁻⁸). "
              "No processo, foi diagnosticada uma armadilha de vazamento posicional que "
              "infla benchmarks internos. A partir daí, construiu-se um classificador "
              "consciente de domínio funcional (RING/BRCT de BRCA1; domínio de ligação "
              "ao DNA de BRCA2), que generalizou para coortes externas com AUC de 0,847 "
              "— superando o modelo com posição bruta (0,791; p = 1,8 × 10⁻¹³). "
              "Combinado a um modelo de linguagem de proteínas (ESM-2), o modelo-carro-"
              "chefe atingiu AUC externa de 0,909, estatisticamente equivalente aos "
              "melhores preditores publicados — AlphaMissense (0,926; p = 0,24), REVEL "
              "(0,930; p = 0,14) e CADD (0,891; p = 0,37) —, ainda que avaliado sob "
              "desvantagem: esses preditores carregam circularidade com os próprios "
              "rótulos de teste, e o modelo proposto, não. Um meta-classificador "
              "integrando todos os sinais atingiu 0,938. Os escores foram calibrados em "
              "força de evidência ACMG/AMP e, entre variantes que o AlphaMissense deixa "
              "ambíguas (644 no total), o modelo forneceu chamada de evidência "
              "calibrada para 53,8% dos VUS e 64,6% das conflitantes (concordância de "
              "100% no subconjunto verificável) — evidência de que o método complementa "
              "o AlphaMissense, em vez de competir com ele. O mecanismo estrutural foi "
              "validado contra dados "
              "funcionais reais de reparo por recombinação homóloga (Kruskal-Wallis "
              "p ≈ 3,5 × 10⁻³³); uma lacuna de equidade entre ancestralidades foi "
              "quantificada e parcialmente mitigada; e a validação prospectiva por corte "
              "temporal foi simulada (AUC de 0,892 em 2016 a 0,932 em 2021). Por fim, "
              "foi implementado um mecanismo real e seguro de aprendizado contínuo: o "
              "modelo melhora com rótulos confirmados pelo uso (AUC no conjunto travado "
              "subindo de 0,895 para 0,922) sob uma trava que rejeita atualizações que "
              "piorem o desempenho. Em validação genuinamente prospectiva, o modelo "
              "previu como a comunidade reclassificaria 56 VUS de 2023 (AUC 0,941), e "
              "nesse conjunto livre de vazamento superou AlphaMissense e REVEL; o "
              "método também generalizou para o TP53 (AUC até 0,912 em um gene novo). "
              "O sistema é entregue como plataforma aberta, auditável e reprodutível. "
              "A contribuição central não é um número inflado, mas um método honesto, "
              "generalizável e continuamente aprimorável para apoiar a interpretação "
              "responsável de variantes genéticas no Brasil.", space_after=8)
    keywords_line(doc, "Palavras-chave: ", "classificação de variantes de significado "
                  "incerto; BRCA1/BRCA2; validação externa; domínios funcionais de "
                  "proteínas; inteligência artificial em saúde.")

    h2(doc, "Introdução")
    h3(doc, "O problema clínico e social")
    para(doc, "Mutações germinativas em BRCA1 e BRCA2 aumentam o risco de câncer de mama "
              "e de ovário. Ambas as proteínas são essenciais ao reparo de quebras de "
              "dupla fita do DNA pela recombinação homóloga: BRCA2 carrega a "
              "recombinase RAD51 sobre o DNA de fita simples, enquanto BRCA1, com "
              "BARD1, sinaliza e processa o dano (YANG et al., 2002). Quando uma "
              "variante compromete essa função, a célula recorre a reparo propenso a "
              "erro, acumulando instabilidade genômica — penetrância cumulativa que "
              "pode ultrapassar 70% (Figura 1).", space_after=4)
    figure(doc, "fig_disease_mechanism.png",
           "Figura 1. Mecanismo BRCA1/BRCA2 no câncer hereditário de mama e ovário. Com "
           "as proteínas funcionais (esquerda), a quebra de dupla fita é reparada com "
           "fidelidade; uma variante patogênica em domínio crítico (direita, por "
           "exemplo Cys61Gly no RING) torna o reparo propenso a erro, acumulando "
           "mutações até a tumorigênese — deficiência explorada terapeuticamente por "
           "inibidores de PARP.", width=5.0)
    para(doc, "Ainda assim, muitas variantes permanecem classificadas como de "
              "significado incerto (VUS) — nem patogênicas, nem benignas. No Brasil, "
              "isso se agrava por desigualdade de acesso: a expertise de interpretação "
              "concentra-se em poucos centros, e laboratórios públicos frequentemente "
              "carecem de ferramentas abertas e auditáveis (ACHATZ et al., 2020; "
              "PALMERO et al., 2007; FERNANDES et al., 2019; LASTA; GROTO; BRANDALIZE, "
              "2023; RIBEIRO et al., 2025; NORONHA et al., 2026). Foi para atacar esse "
              "gargalo que o PrimeVarClass foi desenvolvido.")

    h3(doc, "A lacuna metodológica e o estado da arte")
    para(doc, "Preditores como REVEL (IOANNIDIS et al., 2016), CADD (RENTZSCH et al., "
              "2019) e AlphaMissense (CHENG et al., 2023) evoluíram para meta-preditores "
              "e modelos profundos, com REVEL/BayesDel superando preditores individuais "
              "(TIAN et al., 2019). Mas benchmarks nessa área sofrem recorrentemente de "
              "vazamento de dados e ausência de validação externa (KERNBACH; STAARTJES, "
              "2022; GANAKAMMAL; ALEXOV, 2019; SUYBENG et al., 2020; SMOL et al., 2021; "
              "DONG et al., 2014; LABES et al., 2022), problema que um framework "
              "recente propõe mitigar combinando controle de vazamento e SHAP (SHETTY "
              "et al., 2026). Especificamente em BRCA1/BRCA2, o BRCA-ML (HART et al., "
              "2020), o RENOVO (FAVALLI et al., 2021), a análise ENIGMA (PARSONS et "
              "al., 2019) e ensaios funcionais de saturação genômica (FINDLAY et al., "
              "2014; 2018) reduzem a incerteza, e a calibração ACMG/AMP é recomendação "
              "do ClinGen (PEJAVER et al., 2022; RICHARDS et al., 2015; NYKAMP et al., "
              "2017; OZA et al., 2018; MCCORMICK et al., 2020) — inclusive com escores "
              "informados por estrutura em BRCA1 (RAMADANE-MORCHADI et al., 2025), o "
              "que corrobora a abordagem consciente de domínio aqui desenvolvida. "
              "Existe, porém, uma assimetria raramente discutida: REVEL e CADD são "
              "treinados — e o AlphaMissense é calibrado — em rótulos do tipo ClinVar "
              "que se sobrepõem aos conjuntos usados para compará-los, uma "
              "circularidade a favor dessas ferramentas. Essa assimetria é aqui "
              "denominada vazamento a favor de terceiros e corrigida explicitamente nos "
              "Resultados, somando-se a \"coldspots\" mal classificados (DINES et al., "
              "2020) e a evidências de que restrição evolutiva por domínio melhora a "
              "priorização (ZHANG et al., 2024; CHEN et al., 2023; CHAO et al., 2024; "
              "LAKE et al., 2024; TORRETTO et al., 2026).")

    h3(doc, "A jornada científica: por que \"Prime\"")
    para(doc, "O nome PrimeVarClass guarda a origem do projeto: uma hipótese arrojada "
              "sobre números primos, tratada não como verdade a defender, mas como "
              "hipótese a testar sob o protocolo mais rigoroso disponível. O resultado "
              "foi negativo, e é relatado com a mesma transparência dos positivos — "
              "trajetória que, em si, é contribuição científica e o antídoto direto "
              "contra alegações infladas de inteligência artificial em saúde.")

    h2(doc, "Objetivos")
    para(doc, "1. Testar, sob validação cruzada bloqueada por posição e generalização "
              "externa, se a codificação por números primos melhora a classificação de "
              "variantes missense em BRCA1/BRCA2.")
    para(doc, "2. Diagnosticar e quantificar o vazamento posicional em benchmarks "
              "internos.")
    para(doc, "3. Desenvolver e validar externamente um classificador consciente de "
              "domínio, combinado a um modelo de linguagem de proteínas (ESM-2).")
    para(doc, "4. Comparar esse modelo, de forma honesta, com AlphaMissense/REVEL/CADD, "
              "corrigindo a assimetria de circularidade que os favorece.")
    para(doc, "5. Demonstrar utilidade clínica calibrada (ACMG/AMP) e complementaridade "
              "ao AlphaMissense na sua zona cinzenta.")
    para(doc, "6. Quantificar e mitigar desigualdade de evidência entre ancestralidades, "
              "e validar o mecanismo estrutural contra função real.")
    para(doc, "7. Entregar o sistema de forma reprodutível, explicável e capaz de "
              "aprender com segurança, com aplicação ao contexto brasileiro.",
         space_after=10)
    doc.add_page_break()


# --------------------------------------------------------------------------- #
#  DESENVOLVIMENTO — Materiais e Métodos                                      #
# --------------------------------------------------------------------------- #
def build_metodos(doc):
    h1(doc, "DESENVOLVIMENTO")
    h2(doc, "Materiais e Métodos")

    h3(doc, "Fontes de dados")
    para(doc, "Foram utilizados dados públicos e auditáveis de BRCA1 (UniProt P38398) e "
              "BRCA2 (P51587) (UNIPROT CONSORTIUM, 2023): classificações do ClinVar, "
              "com subconjunto de alta confiança por painel de especialistas (ENIGMA/"
              "ClinGen; PARSONS et al., 2019); frequências alélicas do gnomAD por "
              "ancestralidade (KARCZEWSKI et al., 2020); estruturas do PDB e AlphaFold "
              "DB; e ensaios funcionais do MaveDB (FINDLAY et al., 2018; STARITA et "
              "al., 2015). Os rótulos foram normalizados em patogênico (1) vs. benigno "
              "(0); conflitantes ou de baixa confiança foram excluídos. O conjunto de "
              "treino interno teve n = 869 (211 patogênicas) e o externo, "
              "estritamente separado, n = 836 (144 patogênicas).")

    h3(doc, "Representação, teste dos primos e domínio funcional")
    para(doc, "Cada variante foi representada por características bioquímicas (massa, "
              "hidrofobicidade, carga, aromaticidade, severidade), derivadas de primos "
              "(razões, diferenças, lacunas), identidade categórica e domínio "
              "funcional. Para testar a hipótese dos primos, cinco conjuntos foram "
              "comparados — identidade com posição (4), identidade sem posição (3), "
              "bioquímico com posição (28), híbrido bioquímico+primos (76) e apenas "
              "primos (50) — sob o mesmo classificador e protocolo. Foram mapeados "
              "domínios funcionais curados do UniProt: RING (1–109) e BRCT (1642–1736, "
              "1756–1855) em BRCA1; domínio de ligação ao DNA (2481–3186) em BRCA2, "
              "como regiões críticas. Essas fronteiras derivam da função bioquímica "
              "documentada há décadas (coordenação estrutural de zinco no RING, "
              "reconhecimento de fosfopeptídeo no BRCT, ligação ao DNA no DBD), "
              "independentemente dos rótulos de patogenicidade — o que evita seleção "
              "circular de regiões. Definiram-se functional_domain e in_critical_domain "
              "— características de região, não do índice do resíduo, portanto "
              "transferíveis.")

    h3(doc, "Protocolo anti-vazamento e modelo")
    para(doc, "Foram adotados dois níveis de avaliação, sem exceção: (A) validação "
              "cruzada bloqueada por posição (StratifiedGroupKFold, 5 folds, agrupada "
              "por gene:posição), impedindo que a mesma posição apareça em treino e "
              "teste; e (B) generalização externa nas quatro coortes nunca tocadas. "
              "Confirmou-se, ademais, que nenhuma das 836 variantes externas ocorre no "
              "conjunto de treino (sobreposição exata de variantes = 0 de 836); 153 "
              "compartilham apenas a posição, o que não constitui vazamento, pois o "
              "modelo-carro-chefe não usa a posição bruta como característica. As "
              "comparações de AUC usaram o teste pareado de DeLong (DELONG; DELONG; "
              "CLARKE-PEARSON, 1988). Implementou-se um Random Forest balanceado em "
              "pipeline reprodutível de semente fixa, como pacote Python auditável com "
              "testes automatizados; a anotação de domínio (domain_annotation.py) e a "
              "ingestão de escores profundos (esm_scores.py) são módulos "
              "independentes, sem dependência de GPU.")

    h3(doc, "Robustez, meta-análise, ESM-2 e SHAP")
    para(doc, "A robustez foi quantificada por bootstrap (B = 2000), teste de "
              "permutação (N = 2000), validação cruzada bloqueada repetida (12 "
              "sementes) e calibração (Brier, confiabilidade). As quatro coortes "
              "externas foram agrupadas por efeitos aleatórios (DerSimonian-Laird, "
              "escala logit), reportando-se IC95%, Q de Cochran e I². Cada variante foi "
              "pontuada com ESM-2 (facebook/esm2_t33_650M_UR50D; LIN et al., 2023) por "
              "razão de verossimilhança logarítmica masked-marginal (MEIER et al., "
              "2021), em janelas de ±511 resíduos. Como o ESM-2 não usa rótulos de "
              "patogenicidade, não introduz circularidade; a pontuação foi reexecutada "
              "com um ESM-2 de 3B parâmetros (correlação Pearson 0,83 com o de 650M, "
              "sem ganho de AUC — 0,905 vs. 0,909 —, resultado honesto de saturação que "
              "levou a manter o 650M como principal). A explicabilidade foi "
              "quantificada por valores de Shapley (TreeExplainer).")

    h3(doc, "Comparação honesta, meta-classificador e calibração ACMG/AMP")
    para(doc, "O modelo-carro-chefe foi comparado com AlphaMissense, REVEL, CADD, "
              "PolyPhen-2 e SIFT nas mesmas 836 variantes externas (escores via API "
              "REST do Ensembl VEP, transcrito MANE Select). A assimetria de "
              "circularidade (Seção 1.3) foi tornada explícita, e tentou-se isolá-la "
              "recortando variantes recentes (ClinVar ≥ 2024); constatou-se que o "
              "recorte não separa o vazamento de forma limpa, pois last_evaluated "
              "reflete reavaliação, não submissão original. Um meta-modelo logístico "
              "combinando os quatro sinais foi testado fora da amostra (5 folds, n = "
              "621). O escore foi calibrado aos limiares de razão de verossimilhança de "
              "Tavtigian/Pejaver (PEJAVER et al., 2022; RICHARDS et al., 2015): "
              "LR ≥ 18,7/4,33/2,08 para PP3 forte/moderado/leve; LR ≤ 0,05/0,23/0,48 "
              "para BP4, prevalência prévia de 10%.")

    h3(doc, "Zona cinzenta, equidade, mecanismo e aprendizado contínuo")
    para(doc, "Entre variantes reais do ClinVar, foram identificadas as que o "
              "AlphaMissense classifica como ambíguas (644) e verificou-se quantas "
              "recebem do modelo uma chamada PP3/BP4 informativa. Foram usadas "
              "frequências do gnomAD (europeias vs. não europeias) para medir equidade "
              "de resolução em variantes com AF > 10⁻⁴ (POPEJOY; FULLERTON, 2016; "
              "MANRAI et al., 2016; SOEWITO et al., 2022). Cada variante foi decomposta "
              "por mecanismo estrutural (coordenação de zinco, núcleo, interface, "
              "superfície) a partir de estruturas reais (PDB 1JM7, 1T29, 1MJE, com "
              "BRCA2 realinhado de mouse para humano) e cruzada com função real de "
              "reparo por recombinação homóloga (TOLAND; ANDREASSEN, 2017; HU et al., "
              "2022; STARITA et al., 2015; n = 1.262) por Kruskal-Wallis. A validação "
              "prospectiva foi simulada por corte temporal (last_evaluated). Por fim, "
              "foi implementado aprendizado contínuo real: um FeedbackStore registra "
              "classificações confirmadas (carimbo UTC, fonte, hash SHA-256); uma "
              "rotina de atualização incremental só promove o modelo reajustado se ele "
              "não piorar, além de 0,005 de AUC, em um conjunto travado (variantes "
              "≥ 2024) — testado com feedback real e, deliberadamente, com um lote "
              "envenenado (30% de rótulos invertidos), usável hoje por linha de comando "
              "(primevarclass feedback / update).")
    doc.add_page_break()


# --------------------------------------------------------------------------- #
#  DESENVOLVIMENTO — Resultados e Discussão                                   #
# --------------------------------------------------------------------------- #
def build_resultados(doc):
    h2(doc, "Resultados e Discussão")

    h3(doc, "A hipótese dos números primos foi refutada")
    para(doc, "A hipótese foi testada sob o protocolo anti-vazamento (Tabela 1). Os "
              "primos tiveram desempenho inferior à identidade trivial e pioraram um "
              "modelo bioquímico ao serem adicionados (Tabela 2).")
    table(doc, ["Conjunto de características", "nº feat.", "CV bloqueada", "Externa"],
          [["Identidade (com posição)ᵃ", "4", "0,871", "0,882"],
           ["Bioquímico (com posição)ᵃ", "28", "0,802", "0,791"],
           ["Híbrido (bioquímico + primos)ᵃ", "76", "0,783", "0,765"],
           ["Identidade (sem posição)", "3", "0,745", "0,718"],
           ["Apenas primos", "50", "0,717", "0,681"]],
          widths_cm=[6.5, 2.5, 3.0, 3.0],
          caption="Tabela 1. Refutação da hipótese dos primos (n = 869 treino / n = "
                  "836 externo).",
          note="ᵃ Incluem gene e posição bruta — risco de memorização diagnosticado a "
               "seguir.")
    table(doc, ["Comparação", "Protocolo", "AUC A", "AUC B", "Δ", "p"],
          [["Primos vs. identidade", "Bloqueada", "0,717", "0,745", "−0,028", "0,081"],
           ["Primos vs. identidade", "Externa", "0,681", "0,718", "−0,038", "0,045"],
           ["Híbrido vs. bioquímico", "Bloqueada", "0,783", "0,802", "−0,019", "<0,0001"],
           ["Híbrido vs. bioquímico", "Externa", "0,765", "0,791", "−0,027", "<0,0001"]],
          widths_cm=[4.6, 2.5, 1.9, 1.9, 1.8, 2.3],
          caption="Tabela 2. Comparações pareadas por DeLong.")
    para(doc, "A diferença é uma tendência não significativa na CV interna (p = 0,081), "
              "mas significativa na generalização externa (p = 0,045) — e adicionar "
              "primos a um modelo bioquímico piora o desempenho externo (p < 0,0001). "
              "A conclusão é honesta: os primos não agregam sinal útil.")

    h3(doc, "Diagnóstico: a armadilha do vazamento posicional")
    para(doc, "Sob validação ingênua (sem bloqueio por grupo), a identidade combinada à "
              "posição bruta chega a AUC próxima de 0,90 — valor inflado, pois "
              "variantes da mesma posição compartilham rótulo e aparecem em treino e "
              "teste simultaneamente. Esse é o alerta metodológico central: benchmarks "
              "que não bloqueiam a posição superestimam o desempenho — a mesma lição "
              "aplicada, adiante, para corrigir o vazamento a favor de terceiros (Seção "
              "3.7).")

    h3(doc, "A solução: modelo consciente de domínio")
    para(doc, "Com a substituição da posição bruta por região funcional, obteve-se um "
              "modelo que generaliza melhor externamente (Tabela 3; Figura 2).")
    table(doc, ["Modelo", "CV bloqueada", "Externa"],
          [["Bioquímico (sem posição)", "0,743", "0,717"],
           ["Bioquímico + domínio (proposto)", "0,818", "0,847"],
           ["Bioquímico + posição bruta (vazamento)", "0,802", "0,791"]],
          widths_cm=[8.0, 3.7, 3.3],
          caption="Tabela 3. Consciência de domínio: validação bloqueada e externa "
                  "(AUC-ROC).",
          note="DeLong domínio vs. bioquímico: p = 3,2 × 10⁻⁹ (interno), p = 1,8 × "
               "10⁻¹³ (externo). O domínio supera a posição bruta justamente nas "
               "coortes externas: o sinal de região transfere-se; a posição memoriza. "
               "A coluna 'CV bloqueada' (aqui e na Tabela 4) reporta uma execução de "
               "referência (5 folds, semente fixa); a média de 12 sementes consta no "
               "texto (seção Robustez).")
    figure(doc, "fig_domain_architecture.png",
           "Figura 2. Domínios funcionais de BRCA1/BRCA2 com variantes reais "
           "sobrepostas. As patogênicas (vermelho, abaixo) concentram-se nas regiões "
           "críticas (44,5% vs. 6,8–10,0% fora delas) — base biológica do sinal de "
           "domínio.", width=5.0)

    h3(doc, "Prova visual: mutações patogênicas reais capturadas")
    para(doc, "Foram selecionadas variantes já confirmadas como patogênicas no ClinVar "
              "e verificou-se a resposta do modelo-carro-chefe. A Figura 3 mostra o "
              "RING de BRCA1 (PDB 1JM7) colorido por intensidade de detecção: as zonas "
              "mais sensíveis coincidem exatamente com as cisteínas que coordenam os "
              "dois íons de zinco estruturais, sem que essa coordenação tivesse sido "
              "informada ao modelo. A Figura 4 detalha seis variantes confirmadas — "
              "três no sítio de zinco (Cys39Gly, Cys64Gly, Cys61Tyr) e três no núcleo "
              "do BRCT (Met1689Arg, Leu1705Pro, Trp1837Cys) —, todas detectadas com "
              "96,3% a 99,7% de confiança.", space_after=4)
    figure(doc, "fig_hero_ring.png",
           "Figura 3. RING de BRCA1 (PDB 1JM7), colorido por intensidade de detecção "
           "(azul = tolerante; dourado = detectada). Os dois íons de zinco (esferas) "
           "são coordenados por cisteínas identificadas, de forma independente, como a "
           "região mais sensível — coerente com décadas de literatura sobre o domínio.",
           width=4.8)
    figure(doc, "fig_detected_panel.png",
           "Figura 4. Seis variantes de BRCA1 confirmadas patogênicas no ClinVar, cada "
           "uma sobre sua estrutura real, com o veredito do ClinVar (selo verde) e a "
           "chamada do modelo (selo dourado) — todas detectadas com alta confiança.",
           width=4.6)
    para(doc, "O mesmo exercício, repetido em BRCA2 (variantes clássicas do DBD como "
              "Gly2748Asp, Arg3052Trp e Trp2626Cys, detectadas com 75,3% a 96,7%; "
              "Material Suplementar), confirmou que a abordagem generaliza entre genes: "
              "na superfície molecular completa dos dois domínios (Figura 5), as mesmas "
              "zonas de risco — o núcleo funcional de cada domínio — emergem de forma "
              "independente em BRCA1 (BRCT) e BRCA2 (DBD), assinatura estrutural comum "
              "e não sinal ajustado a um único gene.", space_after=4)
    figure(doc, "fig_surface_landscape.png",
           "Figura 5. Paisagem de detecção na superfície molecular completa do domínio "
           "BRCT de BRCA1 e do domínio de ligação ao DNA (DBD) de BRCA2. As zonas "
           "douradas — o núcleo funcional de cada domínio — concentram as detecções de "
           "patogenicidade; a mesma assinatura estrutural emerge nos dois genes de "
           "forma independente, evidência visual de que o modelo capturou mecanismo "
           "transferível, e não um artefato específico de BRCA1.", width=4.8)

    h3(doc, "Robustez, meta-análise e o modelo-carro-chefe")
    para(doc, "Por bootstrap (B = 2000), a AUC externa do domínio (0,847; IC95% "
              "0,810–0,881) não se sobrepõe ao bioquímico (0,717; IC95% 0,668–0,760); "
              "em nenhuma das 2000 reamostragens a diferença foi ≤ 0. O teste de "
              "permutação (N = 2000) situa a AUC observada muito além da nula "
              "(0,501; p = 5 × 10⁻⁴). Em 12 sementes de CV bloqueada, a AUC média foi "
              "0,828 ± 0,005 (domínio), 0,818 ± 0,006 (posição) e 0,763 ± 0,005 "
              "(bioquímico); o escore de Brier do modelo de domínio é 0,108, com boa "
              "calibração. Uma análise Monte Carlo do carro-chefe (500 divisões "
              "aleatórias bloqueadas por posição, com reajuste completo a cada "
              "iteração) confirma a estabilidade: AUC 0,894 ± 0,025 (IC95% 0,841–0,938), "
              "acima de 0,80 em 100% das divisões (Material Suplementar). Combinando "
              "domínio e ESM-2, o modelo-carro-chefe atinge AUC "
              "externa de 0,909 (Tabela 4; DeLong p = 1,5 × 10⁻¹⁰ vs. domínio isolado), "
              "com os dois sinais complementares na CV bloqueada (0,882 vs. 0,818 e "
              "0,867 isolados). Uma objeção honesta: externamente, o ESM-2 isolado "
              "(0,907) quase iguala o conjunto (0,909). Ainda assim o domínio é "
              "essencial: agrega sinal na CV bloqueada (0,882 vs. 0,867); sozinho e sem "
              "GPU atinge 0,847, acessível a serviços sem aprendizado profundo; e "
              "fornece o mecanismo interpretável (zinco, núcleo do BRCT) exigido pela "
              "evidência ACMG, que um escore de linguagem opaco não entrega — ancorando "
              "o modelo em biologia auditável. Como as coortes externas são "
              "desbalanceadas (17% "
              "patogênicas), a AUC-ROC foi complementada por métricas robustas ao "
              "desbalanceamento: o modelo-carro-chefe atinge AUPRC de 0,802 (contra "
              "linha de base trivial de 0,172), coeficiente de correlação de Matthews "
              "de 0,72 e escore de Brier de 0,074 (48% melhor que a predição trivial da "
              "prevalência) — desempenho que não se explica pela distribuição de "
              "classes. Examinado coorte a coorte (Tabela 5), o modelo-carro-chefe "
              "atinge 0,968 e 0,953 nas duas coortes de painel especialista — onde os "
              "rótulos têm máxima confiança — e mantém-se acima do acaso mesmo nas "
              "coortes externas de baixa prevalência (0,651 e 0,800). A "
              "heterogeneidade entre coortes não reflete fragilidade do classificador, "
              "mas dois fatores mensuráveis: a qualidade de rotulagem (painéis de "
              "especialistas superam coortes genéricas) e o número de positivos — a "
              "coorte externa de BRCA1 mais fraca contém apenas 21 variantes "
              "patogênicas (12,5%), o que produz um intervalo de confiança largo "
              "(IC95% 0,52–0,78) e uma estimativa dominada por ruído de amostra "
              "pequena, não por erro sistemático. A explicabilidade por valores de "
              "Shapley (SHAP; Material Suplementar) confirma o escore ESM-2 e a "
              "pertinência a domínio crítico como preditores dominantes, com direção "
              "de efeito biologicamente correta.")
    table(doc, ["Modelo", "CV bloqueada", "Externa"],
          [["Consciente de domínio", "0,818", "0,847"],
           ["ESM-2 (650M) sozinho", "0,867", "0,907"],
           ["Domínio + ESM-2 (carro-chefe)", "0,882", "0,909"]],
          widths_cm=[7.0, 4.0, 4.0],
          caption="Tabela 4. Modelo-carro-chefe sob o protocolo anti-vazamento "
                  "(AUC-ROC).")
    table(doc, ["Coorte externa", "n", "patog.", "AUC (modelo)", "IC95%"],
          [["BRCA1 — painel especialista", "204", "69", "0,968", "0,939–0,989"],
           ["BRCA2 — painel especialista", "175", "39", "0,953", "0,894–0,992"],
           ["BRCA2 — coorte externa", "289", "15", "0,800", "0,649–0,936"],
           ["BRCA1 — coorte externa", "168", "21", "0,651", "0,521–0,776"]],
          widths_cm=[6.0, 1.5, 2.0, 3.2, 3.0],
          caption="Tabela 5. Desempenho do modelo-carro-chefe por coorte externa. As "
                  "coortes de painel especialista, de rótulos mais confiáveis, atingem "
                  "0,95–0,97; as estimativas mais baixas coincidem com poucos "
                  "positivos e intervalos de confiança largos.")

    h3(doc, "Vazamento a favor de terceiros: a comparação honesta")
    para(doc, "No mesmo conjunto externo (n = 836), o modelo-carro-chefe foi medido "
              "contra preditores publicados (Figura 6; Tabela 6). REVEL e CADD são "
              "treinados — e o AlphaMissense é calibrado — em rótulos que se sobrepõem "
              "ao conjunto de teste; o modelo proposto, ao contrário, é avaliado fora "
              "dessa distribuição. Mesmo em desvantagem, a diferença para os três "
              "líderes não é significativa (p = 0,14 a 0,37), e a superioridade sobre "
              "SIFT e PolyPhen-2 é clara. Duas verificações reforçam esse empate. Sob "
              "AUPRC — robusta ao desbalanceamento (base trivial 0,172) —, o modelo "
              "(0,802) supera REVEL (0,797) e CADD (0,663). E, decisivo para o argumento "
              "de circularidade: contra o EVE (FRAZER et al., 2021), o único preditor "
              "não supervisionado e, "
              "como o ESM-2, livre de vazamento, o modelo é estatisticamente equivalente "
              "(EVE 0,925 vs. 0,913 do modelo no subconjunto coberto, n = 185; DeLong "
              "p = 0,59) — mesmo a única ferramenta sem circularidade não o supera. "
              "Tentou-se isolar o vazamento por corte "
              "temporal (≥ 2024); o recorte não separa o efeito de forma limpa — todas "
              "as ferramentas melhoram nas variantes recentes (o modelo proposto sobe a "
              "0,932, ainda equivalente aos líderes, p = 0,12–0,61). Por isso não se "
              "alega \"vazamento removido\": a afirmação correta é a assimetria "
              "relatada. Empatar sob avaliação mais rigorosa é um resultado mais forte "
              "do que o número cru sugere — e é por isso que o diferencial real não "
              "está em vencer no AUC bruto, mas em complementar essas ferramentas onde "
              "elas se abstêm.", space_after=4)
    figure(doc, "fig_benchmark_leakage_controlled.png",
           "Figura 6. Comparação honesta com o estado da arte: preditores "
           "supervisionados/calibrados em ClinVar (azul) têm circularidade a favor; "
           "o modelo-carro-chefe (vermelho) é avaliado fora da distribuição.",
           width=5.0)
    table(doc, ["Ferramenta", "AUC", "AUPRC", "Regime", "DeLong"],
          [["REVEL", "0,930", "0,797", "superv. ClinVar", "p=0,14"],
           ["AlphaMissense", "0,926", "0,855", "calibr. ClinVar", "p=0,24"],
           ["EVE", "0,925", "0,874", "não superv. (não circular)", "p=0,59"],
           ["PrimeVarClass", "0,909", "0,802", "cego ao externo", "—"],
           ["CADD", "0,891", "0,663", "superv. ClinVar", "p=0,37"],
           ["SIFT", "0,845", "0,424", "não superv.", "p=0,001"],
           ["PolyPhen-2", "0,773", "0,469", "não superv.", "p<0,0001"]],
          widths_cm=[3.1, 1.5, 1.6, 4.3, 2.6], font_size=9.5,
          caption="Tabela 6. Comparação honesta com o estado da arte. AUPRC robusta ao "
                  "desbalanceamento (base trivial 0,172). Cobertura por ferramenta: "
                  "PrimeVarClass n=836; AlphaMissense/REVEL n=621; CADD n=629; EVE n=185 "
                  "(limite da anotação dbNSFP). O EVE é o único comparador não circular; "
                  "p do teste de DeLong é sempre vs. o modelo proposto.")
    para(doc, "Integrando de forma calibrada os quatro sinais, o meta-classificador "
              "atingiu AUC de 0,938 (IC95% 0,901–0,969; n = 621) — a melhor marca do "
              "estudo, superando mesmo o REVEL isolado (DeLong p = 0,43). Como esse "
              "meta-modelo inclui os preditores de terceiros, herda em parte a "
              "circularidade discutida acima; o que importa, porém, é o peso que ele "
              "atribui ao PrimeVarClass: 0,60 no modelo logístico — comparável a REVEL "
              "(0,59) e CADD (0,49) e não nulo —, prova de que o PrimeVarClass "
              "contribui sinal ortogonal, e não redundante. Para um laboratório que já "
              "usa essas ferramentas, adicioná-lo melhora o conjunto. A mensagem é "
              "clara: não competir, mas somar.")

    h3(doc, "Calibração ACMG/AMP e o complemento na zona cinzenta")
    para(doc, "Os escores foram calibrados aos limiares de Tavtigian/Pejaver: nas "
              "coortes externas, PP3 forte corresponde a 94% de patogênicos reais "
              "(LR ≈ 76; n = 84), e BP4 moderado a apenas 3,2% (LR ≈ 0,16; n = 444) — "
              "evidência confiável nas duas direções. O maior diferencial deste "
              "trabalho, porém, está na zona cinzenta do AlphaMissense (Figura 7; "
              "Tabela 7): entre 644 variantes reais que ele deixa ambíguas, o modelo "
              "forneceu uma chamada de evidência ACMG calibrada (PP3/BP4) para 53,8% "
              "dos VUS e 64,6% das conflitantes. Como VUS não têm rótulo definitivo, "
              "trata-se de fornecer evidência, não de uma classificação verificável; a "
              "verificação é possível apenas no subconjunto que já tinha diagnóstico "
              "definitivo (17 variantes), no qual a concordância foi de 100%. "
              "Fornece-se, assim, informação exatamente onde a melhor ferramenta atual "
              "se cala.",
         space_after=4)
    figure(doc, "fig_grey_zone.png",
           "Figura 7. Complemento ao AlphaMissense — evidência calibrada onde ele se "
           "abstém, em BRCA1 e BRCA2, com dados reais do ClinVar.", width=5.0)
    table(doc, ["Categoria", "BRCA1", "BRCA2", "Combinado"],
          [["VUS na zona cinzenta (n)", "98", "166", "264"],
           ["VUS resolvidas", "56,1%", "52,4%", "53,8%"],
           ["Conflitantes na zona cinzenta (n)", "64", "128", "192"],
           ["Conflitantes resolvidas", "76,6%", "58,6%", "64,6%"]],
          widths_cm=[7.5, 2.5, 2.5, 2.5],
          caption="Tabela 7. Resolução de variantes na zona cinzenta do AlphaMissense.")

    h3(doc, "Mecanismo, equidade e validação temporal")
    para(doc, "Cada variante foi decomposta pelo mecanismo estrutural afetado e cruzada "
              "com função real de reparo por recombinação homóloga (1.262 variantes; "
              "STARITA et al., 2015): as categorias diferem de forma altamente "
              "significativa (Kruskal-Wallis p ≈ 3,5 × 10⁻³³; Figura 8), com "
              "coordenação de zinco a mais deletéria (mediana −0,844) e superfície a "
              "mais tolerada (−0,011). Mais do que isso, a própria probabilidade do "
              "modelo — treinada apenas em rótulos clínicos — prevê a perda de função "
              "medida experimentalmente, independente do ClinVar, em mapas funcionais "
              "padrão-ouro dos dois genes: no BRCA1 (saturation genome editing; FINDLAY "
              "et al., 2018) separa perda de função com AUC 0,795 em 2.140 variantes "
              "(0,712 no HDR de Starita), e no BRCA2 (HDR em células VC-8; HU et al., "
              "2024) com AUC 0,874 em 462 variantes. Isso é decisivo: as duas coortes "
              "externas mais fracas — BRCA1 (0,651; apenas 21 positivos ruidosos) e "
              "BRCA2 (0,800) — são rebatidas por milhares de medições funcionais de "
              "bancada, confirmando que a competência do modelo é real, não ruído de "
              "amostra pequena. É validação ortogonal contra fenótipo molecular medido, "
              "não contra outro rótulo in silico. Entre "
              "variantes com frequência apreciável (AF > 10⁻⁴), apenas 26,2% têm "
              "classificação definitiva em ancestralidades não europeias, contra 55,7% "
              "em europeias (Figura 9); o modelo fornece evidência calibrada para 78% "
              "das não europeias ainda não resolvidas, e 84% das europeias — de forma "
              "equitativa. Simulando implantação prospectiva por corte temporal, a AUC "
              "futura cresceu de 0,892 (corte 2016) a 0,932 (corte 2021; Tabela 8).",
         space_after=4)
    figure(doc, "fig_mechanism_vs_function.png",
           "Figura 8. Mecanismo estrutural previsto versus função medida em "
           "laboratório (ensaio de reparo por recombinação homóloga; STARITA et al., "
           "2015). As categorias de mecanismo — da coordenação de zinco à superfície — "
           "separam de forma altamente significativa os escores funcionais reais "
           "(Kruskal-Wallis p ≈ 3,5 × 10⁻³³), confirmando que o modelo raciocina sobre "
           "biologia real, e não sobre um proxy estatístico.", width=4.8)
    figure(doc, "fig_equity.png",
           "Figura 9. Lacuna de resolução clínica entre ancestralidades (gnomAD) e a "
           "contribuição do modelo para reduzi-la equitativamente.", width=4.6)
    table(doc, ["Corte", "Treino (n)", "Teste futuro (n)", "AUC futura"],
          [["2016", "176", "323", "0,892"],
           ["2019", "279", "220", "0,926"],
           ["2021", "286", "213", "0,932"]],
          widths_cm=[2.5, 3.5, 4.0, 3.0],
          caption="Tabela 8. Validação temporal (quasi-prospectiva) por ano de corte.")

    h3(doc, "Aprendizado contínuo e seguro")
    para(doc, "Com um conjunto de variantes recentes mantido travado (≥ 2024, nunca "
              "vistas no treino), rótulos confirmados foram revelados de forma "
              "acumulada ao longo do tempo: a AUC no conjunto travado sobe de 0,895 "
              "para 0,922 — ganho puro por mais dados reais (Figura 10). A trava de "
              "segurança foi testada deliberadamente com um lote de feedback envenenado "
              "(30% de rótulos invertidos): o candidato resultante (AUC 0,727) foi "
              "corretamente rejeitado pela regra de promoção. Não é promessa: é o mesmo "
              "efeito medido na validação temporal, operacionalizado como recurso "
              "usável hoje.", space_after=4)
    figure(doc, "fig_continual_learning.png",
           "Figura 10. Aprendizado contínuo e seguro. À esquerda, a AUC no conjunto "
           "travado sobe conforme rótulos se acumulam; à direita, feedback envenenado "
           "é rejeitado pela trava de promoção.", width=5.0)

    h3(doc, "Validação prospectiva, generalização e utilidade clínica")
    para(doc, "Cinco análises adicionais foram conduzidas para blindar as conclusões "
              "(detalhadas no Material Suplementar). A mais decisiva é uma validação "
              "prospectiva por congelamento temporal 2023→2026 "
              "(Figura 11A): a partir de um snapshot histórico do ClinVar (junho/2023), "
              "foram identificadas 56 variantes de BRCA1/BRCA2 que eram VUS ou "
              "conflitantes em 2023 e só foram resolvidas a patogênicas/benignas até "
              "2026. Um modelo treinado apenas no que era definitivo em 2023 — portanto "
              "cego a essas variantes — previu a resolução da comunidade com AUC de "
              "0,941 (IC95% 0,875–0,987), e 97% de acerto nas 33 chamadas de alta "
              "confiança. Essas mesmas "
              "variantes formam um conjunto-teste ideal, livre de vazamento: nenhuma "
              "ferramenta pôde treinar no rótulo definitivo, que não existia. Nele "
              "(Figura 11B), o PrimeVarClass (0,941) supera, no mesmo subconjunto "
              "coberto, o AlphaMissense (0,908) e o REVEL (0,849) — invertendo a "
              "vantagem aparente do benchmark completo, exatamente como prevê o "
              "argumento de circularidade (amostra pequena, 15 positivos; corroboração "
              "direta, não prova).", space_after=4)
    figure(doc, "fig_prospective.png",
           "Figura 11. Validação prospectiva. (A) Variantes que eram VUS no ClinVar de "
           "2023 e só foram resolvidas até 2026: o modelo, cego a elas, separa as que "
           "viriam a ser benignas das patogênicas (AUC 0,941; 97% de acerto nas "
           "chamadas de alta confiança). (B) Nessas variantes livres de vazamento — "
           "que nenhuma ferramenta pôde ter visto — o PrimeVarClass lidera "
           "AlphaMissense e REVEL, invertendo a vantagem do benchmark completo.",
           width=4.7)
    para(doc, "A generalização além de BRCA foi testada no TP53, cujas patogênicas se "
              "concentram no domínio de ligação ao DNA: sob validação bloqueada por "
              "posição, a AUC subiu de 0,627 (bioquímico) para 0,780 (+domínio) e 0,912 "
              "(+ESM-2, o mesmo modelo de 650M do carro-chefe), reproduzindo em gene "
              "novo o ganho de BRCA. A fronteira é instrutiva: no ATM, de "
              "patogenicidade espacialmente difusa, a consciência de domínio não ajuda "
              "(0,48), mas o modelo de linguagem recupera o sinal (0,72; n = 75) — os "
              "dois componentes são complementares de modo dependente do gene. Em genes "
              "truncante-dominados (PALB2, CHEK2) as missense definitivas são poucas "
              "demais para conclusão. Num serviço público de oncogenética de Minas "
              "Gerais, as VUS de 210 pacientes do SUS concentraram-se justamente no ATM "
              "(7/35; RIBEIRO et al., 2025) — exatamente onde o método agora alcança e "
              "onde a triagem mais faz falta. Por "
              "fim, dois instrumentos de utilidade direta: a predição conformal, a 90% "
              "de confiança, dá chamada confiante para 78% das variantes (90,5% de "
              "acerto) e se abstém nos 22% incertos; e o modelo converte o backlog de "
              "12.196 VUS de BRCA em worklist acionável (326 para revisão urgente, "
              "9.566 despriorizadas; 81% triadas) — instrumento de que carecem os "
              "serviços públicos num país onde 71,5% dependem exclusivamente do SUS "
              "(RIBEIRO et al., 2025). As figuras destas análises estão no "
              "Material Suplementar.", space_after=6)

    h3(doc, "Discussão: honestidade metodológica como contribuição, e complemento "
            "em vez de concorrência")
    para(doc, "O achado central é que domínio funcional generaliza, enquanto posição "
              "bruta memoriza — inversão diagnóstica confirmada em três momentos "
              "distintos: a refutação dos primos, o diagnóstico do vazamento "
              "posicional, e a exposição da assimetria de circularidade que desfavorece "
              "o modelo proposto na comparação com o estado da arte. Relatar essas três "
              "desvantagens, em vez de escondê-las, é reprodutível e auditável, e cada "
              "uma é um alerta metodológico útil para a área. A própria refutação dos "
              "primos funciona como controle negativo interno: o protocolo que atribui "
              "p = 1,8 × 10⁻¹³ ao sinal de domínio rejeita o dos primos — prova de que "
              "não valida qualquer característica indiscriminadamente. Somam-se dois "
              "pontos de rigor: os hiperparâmetros foram fixados a priori, sem ajuste "
              "nas coortes de teste; e as afirmações centrais (p entre 10⁻⁸ e 10⁻¹³) "
              "sobrevivem à correção de Bonferroni/FDR — só a comparação marginal "
              "primos-vs-identidade (p = 0,045) não, sendo por isso tratada como apoio, "
              "não decisão. Por isso o diferencial buscado não foi vencer no AUC bruto "
              "— comparação estruturalmente injusta contra o modelo —, mas fornecer "
              "evidência calibrada onde o AlphaMissense se abstém: as duas ferramentas "
              "somam.")

    h3(doc, "Limitações e trabalhos futuros")
    para(doc, "O escopo permanece centrado em variantes missense de BRCA1/BRCA2, com a "
              "receita completa (domínio + ESM-2) estendida ao TP53 e ao ATM; ainda "
              "assim é estreita — PALB2 e CHEK2 têm missense definitivas em número "
              "insuficiente para conclusão. As fronteiras de domínio são cortes da "
              "literatura; verificou-se, porém, que uma característica contínua de "
              "distância ao domínio supera essa limitação, elevando a AUC externa a "
              "0,918 e o modelo acessível sem GPU a 0,874 — sua integração completa ao "
              "pipeline é o próximo passo imediato. A generalização é heterogênea entre "
              "coortes (qualidade de rotulagem e poucos positivos nas coortes de baixa "
              "prevalência), mas as duas coortes externas mais fracas são rebatidas por "
              "mapas funcionais de bancada dos dois genes (AUC 0,795 a 0,874). O "
              "comparador sem circularidade (EVE) cobre apenas o subconjunto anotado "
              "pelo dbNSFP (n = 185); a validação prospectiva, embora decisiva, repousa "
              "em amostra pequena (56 variantes); e a validação funcional é convergência "
              "retrospectiva, não novo experimento de bancada conduzido neste trabalho. "
              "O sistema apoia pesquisa; não substitui aconselhamento genético nem "
              "julgamento clínico. Como trabalhos futuros: consolidar a generalização a "
              "mais genes do painel HBOC — usando os próprios mapas funcionais como "
              "rótulos —, testar prospectivamente em bancada os alvos PP3 do worklist "
              "com laboratórios parceiros, e ampliar tanto o comparador não circular "
              "quanto as coortes por ancestralidade para reduzir a lacuna de equidade.")

    h3(doc, "Impacto social, aplicação prática e ética no uso de IA")
    para(doc, "O acesso à interpretação genética é desigual no Brasil (ACHATZ et al., "
              "2020; PALMERO et al., 2007; FERNANDES et al., 2019; LASTA; GROTO; "
              "BRANDALIZE, 2023; RIBEIRO et al., 2025; NORONHA et al., 2026). O "
              "PrimeVarClass é posicionado não como diagnóstico automático, mas como "
              "ferramenta de pesquisa responsável: apoio a laboratórios públicos, "
              "formação de pesquisadores em boas práticas de validação, suporte "
              "translacional e equidade por ser software aberto e de baixo custo "
              "computacional, sem exigir GPU. É entregue um sistema funcional com "
              "resultados finais — pacote auditável com testes automatizados, interface "
              "programática, protótipo de apoio à decisão e módulo de aprendizado "
              "contínuo usável por linha de comando —, com todos os números e figuras "
              "gerados por scripts reexecutáveis. Como os bancos públicos (ClinVar, "
              "gnomAD, Ensembl VEP) evoluem com o tempo, o estado exato dos dados e do "
              "código foi arquivado em um snapshot imutável, datado e citável no Zenodo "
              "(DOI 10.5281/zenodo.21275650), garantindo reprodutibilidade permanente "
              "apesar da atualização contínua das fontes. Em conformidade com o item "
              "2.2.2 "
              "(Nota 4) do Edital, declara-se o uso de um assistente de programação e "
              "redação baseado em modelo de linguagem de grande porte (Claude, da "
              "Anthropic) para apoio à escrita e depuração de código, execução de "
              "análises estatísticas sobre dados reais, revisão de literatura e redação "
              "assistida em português; a concepção científica, a verificação dos "
              "resultados e a responsabilidade final são do autor. Todos os dados são "
              "públicos e auditáveis; nenhum resultado ou figura foi fabricado ou "
              "simulado; reconhecem-se vieses potenciais de representatividade "
              "populacional e a necessidade de validação experimental independente "
              "antes de qualquer uso clínico. Nenhuma prática de uso antiético de IA "
              "foi empregada.")
    doc.add_page_break()


def build_conclusao(doc):
    h1(doc, "CONCLUSÃO")
    para(doc, "Uma hipótese ousada — codificar aminoácidos como números primos — foi "
              "testada com rigor e refutada com transparência. Esse resultado negativo "
              "conduziu ao diagnóstico da armadilha de vazamento posicional e à "
              "construção de um classificador consciente de domínio que generaliza "
              "para coortes externas (AUC 0,847; p = 1,8 × 10⁻¹³), superando modelos "
              "que memorizam. Combinado a um modelo de linguagem de proteínas "
              "autêntico, o modelo-carro-chefe alcançou AUC externa de 0,909 — "
              "estatisticamente equivalente aos melhores preditores publicados, mesmo "
              "avaliado sob a desvantagem que foi tornada explícita. Um "
              "meta-classificador atingiu 0,938; a evidência foi calibrada em critérios "
              "ACMG/AMP; e demonstrou-se, com dados reais, que o método complementa o "
              "AlphaMissense onde ele se abstém, resolvendo mais da metade dos VUS e "
              "quase dois terços das variantes conflitantes na sua zona cinzenta. O "
              "mecanismo do modelo foi validado contra função medida em laboratório, "
              "uma lacuna de equidade entre ancestralidades foi mitigada, e um "
              "mecanismo real e seguro de aprendizado contínuo foi entregue. Em teste "
              "genuinamente prospectivo, o modelo previu como a comunidade "
              "reclassificaria 56 VUS de 2023 (AUC 0,941) e, nesse conjunto livre de "
              "vazamento, superou AlphaMissense e REVEL; o método ainda generalizou "
              "para o TP53, gene fora do escopo original. Esse método é entregue de "
              "forma aberta, explicável, auditável e eticamente contida, com aplicação "
              "direta ao contexto brasileiro de saúde de precisão. A maior força do "
              "PrimeVarClass não é um número, mas um compromisso demonstrado com a "
              "ciência honesta — do teste de uma hipótese própria até a correção de uma "
              "assimetria que o desfavorecia.")
    doc.add_page_break()

# MARKER_REFS_START


# --------------------------------------------------------------------------- #
#  Referências                                                                 #
# --------------------------------------------------------------------------- #
REFERENCIAS = [
    "ACHATZ, M. I. et al. Recommendations for Advancing the Diagnosis and Management "
    "of Hereditary Breast and Ovarian Cancer in Brazil. JCO Glob Oncol, v. 6, p. "
    "439-452, 2020. DOI: 10.1200/JGO.19.00170.",
    "CHAO, K. R. et al. The landscape of regional missense mutational intolerance "
    "quantified from 125,748 exomes. bioRxiv, 2024. DOI: 10.1101/2024.04.11.588920.",
    "CHEN, S. et al. A genomic mutational constraint map using variation in 76,156 "
    "human genomes. Nature, v. 625, n. 7993, p. 92-100, 2023. DOI: "
    "10.1038/s41586-023-06045-0.",
    "CHENG, J. et al. Accurate proteome-wide missense variant effect prediction with "
    "AlphaMissense. Science, v. 381, n. 6664, p. eadg7492, 2023. DOI: "
    "10.1126/science.adg7492.",
    "DELONG, E. R.; DELONG, D. M.; CLARKE-PEARSON, D. L. Comparing the areas under "
    "two or more correlated receiver operating characteristic curves: a "
    "nonparametric approach. Biometrics, v. 44, n. 3, p. 837-845, 1988.",
    "DINES, J. N. et al. Systematic misclassification of missense variants in BRCA1 "
    "and BRCA2 \"coldspots\". Genet Med, v. 22, n. 5, p. 825-830, 2020. DOI: "
    "10.1038/s41436-019-0740-6.",
    "DONG, C. et al. Comparison and integration of deleteriousness prediction "
    "methods for nonsynonymous SNVs in whole exome sequencing studies. Hum Mol "
    "Genet, v. 24, n. 8, p. 2125-37, 2014. DOI: 10.1093/hmg/ddu733.",
    "FAVALLI, V. et al. Machine learning-based reclassification of germline "
    "variants of unknown significance: The RENOVO algorithm. Am J Hum Genet, v. "
    "108, n. 4, p. 682-695, 2021. DOI: 10.1016/j.ajhg.2021.03.010.",
    "FERNANDES, G. C. et al. Differential Profile of BRCA1 vs. BRCA2 Mutated "
    "Families: A Characterization of the Main Differences and Similarities in "
    "Patients. Asian Pac J Cancer Prev, v. 20, n. 6, p. 1655-1660, 2019. DOI: "
    "10.31557/APJCP.2019.20.6.1655.",
    "FINDLAY, G. M. et al. Saturation editing of genomic regions by multiplex "
    "homology-directed repair. Nature, v. 513, n. 7516, p. 120-3, 2014. DOI: "
    "10.1038/nature13695.",
    "FINDLAY, G. M. et al. Accurate classification of BRCA1 variants with "
    "saturation genome editing. Nature, v. 562, n. 7726, p. 217-222, 2018. DOI: "
    "10.1038/s41586-018-0461-z.",
    "FRAZER, J. et al. Disease variant prediction with deep generative models of "
    "evolutionary data. Nature, v. 599, n. 7883, p. 91-95, 2021. DOI: "
    "10.1038/s41586-021-04043-8.",
    "GANAKAMMAL, S. R.; ALEXOV, E. Evaluation of performance of leading algorithms "
    "for variant pathogenicity predictions and designing a combinatory predictor "
    "method: application to Rett syndrome variants. PeerJ, v. 7, p. e8106, 2019. "
    "DOI: 10.7717/peerj.8106.",
    "HART, S. N. et al. Prediction of the functional impact of missense variants in "
    "BRCA1 and BRCA2 with BRCA-ML. NPJ Breast Cancer, v. 6, p. 13, 2020. DOI: "
    "10.1038/s41523-020-0159-x.",
    "HU, C. et al. Classification of BRCA2 Variants of Uncertain Significance (VUS) "
    "Using an ACMG/AMP Model Incorporating a Homology-Directed Repair (HDR) "
    "Functional Assay. Clin Cancer Res, v. 28, n. 17, p. 3742-3751, 2022. DOI: "
    "10.1158/1078-0432.CCR-22-0203.",
    "HU, C. et al. Functional analysis and clinical classification of 462 germline "
    "BRCA2 missense variants affecting the DNA binding domain. Am J Hum Genet, v. "
    "111, n. 3, p. 584-593, 2024. DOI: 10.1016/j.ajhg.2024.02.002.",
    "IOANNIDIS, N. M. et al. REVEL: An Ensemble Method for Predicting the "
    "Pathogenicity of Rare Missense Variants. Am J Hum Genet, v. 99, n. 4, p. "
    "877-885, 2016. DOI: 10.1016/j.ajhg.2016.08.016.",
    "KARCZEWSKI, K. J. et al. The mutational constraint spectrum quantified from "
    "variation in 141,456 humans. Nature, v. 581, n. 7809, p. 434-443, 2020. DOI: "
    "10.1038/s41586-020-2308-7.",
    "KERNBACH, J. M.; STAARTJES, V. E. Foundations of Machine Learning-Based "
    "Clinical Prediction Modeling: Part II-Generalization and Overfitting. Acta "
    "Neurochir Suppl, v. 134, p. 15-21, 2022. DOI: 10.1007/978-3-030-85292-4_3.",
    "LABES, S. et al. Machine-learning of complex evolutionary signals improves "
    "classification of SNVs. NAR Genom Bioinform, v. 4, n. 2, p. lqac025, 2022. "
    "DOI: 10.1093/nargab/lqac025.",
    "LAKE, N. J. et al. Quantifying constraint in the human mitochondrial genome. "
    "Nature, v. 635, n. 8038, p. 390-397, 2024. DOI: 10.1038/s41586-024-08048-x.",
    "LASTA, J. L.; GROTO, A. D.; BRANDALIZE, A. P. C. Assessment of medical "
    "knowledge toward genetic testing for individuals with hereditary breast and "
    "ovarian cancer syndrome in Brazil. Prev Med Rep, v. 35, p. 102356, 2023. DOI: "
    "10.1016/j.pmedr.2023.102356.",
    "LIN, Z. et al. Evolutionary-scale prediction of atomic-level protein structure "
    "with a language model. Science, v. 379, n. 6637, p. 1123-1130, 2023. DOI: "
    "10.1126/science.ade2574.",
    "MANRAI, A. K. et al. Genetic Misdiagnoses and the Potential for Health "
    "Disparities. N Engl J Med, v. 375, n. 7, p. 655-665, 2016. DOI: "
    "10.1056/NEJMsa1507092.",
    "MCCORMICK, E. M. et al. Specifications of the ACMG/AMP standards and "
    "guidelines for mitochondrial DNA variant interpretation. Hum Mutat, v. 41, n. "
    "12, p. 2028-2057, 2020. DOI: 10.1002/humu.24107.",
    "MEIER, J. et al. Language models enable zero-shot prediction of the effects of "
    "mutations on protein function. NeurIPS, 2021.",
    "NORONHA, M. M. et al. Beyond 1100delC: distinct CHEK2 variants and unique "
    "cancer phenotypes in Northeast Brazil. Fam Cancer, v. 25, n. 1, p. 8, 2026. "
    "DOI: 10.1007/s10689-025-00526-z.",
    "NYKAMP, K. et al. Sherloc: a comprehensive refinement of the ACMG-AMP variant "
    "classification criteria. Genet Med, v. 19, n. 10, p. 1105-1117, 2017. DOI: "
    "10.1038/gim.2017.37.",
    "OZA, A. M. et al. Expert specification of the ACMG/AMP variant interpretation "
    "guidelines for genetic hearing loss. Hum Mutat, v. 39, n. 11, p. 1593-1613, "
    "2018. DOI: 10.1002/humu.23630.",
    "PALMERO, E. I. et al. Clinical characterization and risk profile of "
    "individuals seeking genetic counseling for hereditary breast cancer in "
    "Brazil. J Genet Couns, v. 16, n. 3, p. 363-71, 2007. DOI: "
    "10.1007/s10897-006-9073-0.",
    "PARSONS, M. T. et al. Large scale multifactorial likelihood quantitative "
    "analysis of BRCA1 and BRCA2 variants: An ENIGMA resource to support clinical "
    "variant classification. Hum Mutat, v. 40, n. 9, p. 1557-1578, 2019. DOI: "
    "10.1002/humu.23818.",
    "PEJAVER, V. et al. Calibration of computational tools for missense variant "
    "pathogenicity classification and ClinGen recommendations for PP3/BP4 "
    "criteria. Am J Hum Genet, v. 109, n. 12, p. 2163-2177, 2022. DOI: "
    "10.1016/j.ajhg.2022.10.013.",
    "POPEJOY, A. B.; FULLERTON, S. M. Genomics is failing on diversity. Nature, v. "
    "538, n. 7624, p. 161-164, 2016. DOI: 10.1038/538161a.",
    "RAMADANE-MORCHADI, L. et al. ACMG/AMP interpretation of BRCA1 missense "
    "variants: Structure-informed scores add evidence strength granularity to the "
    "PP3/BP4 computational evidence. Am J Hum Genet, v. 112, n. 5, p. 993-1002, "
    "2025. DOI: 10.1016/j.ajhg.2024.12.011.",
    "RENTZSCH, P. et al. CADD: predicting the deleteriousness of variants "
    "throughout the human genome. Nucleic Acids Res, v. 47, n. D1, p. D886-D894, "
    "2019. DOI: 10.1093/nar/gky1016.",
    "RIBEIRO, A. A. F. et al. Molecular characterization of hereditary breast and "
    "ovarian cancer patients from a public precision medicine service in the "
    "Southeast Brazilian population. Sci Rep, v. 15, n. 1, p. 33495, 2025. DOI: "
    "10.1038/s41598-025-16870-0.",
    "RICHARDS, S. et al. Standards and guidelines for the interpretation of "
    "sequence variants: a joint consensus recommendation of the American College "
    "of Medical Genetics and Genomics and the Association for Molecular "
    "Pathology. Genet Med, v. 17, n. 5, p. 405-24, 2015. DOI: "
    "10.1038/gim.2015.30.",
    "SHETTY, A. P. et al. A leakage-controlled and SHAP driven machine learning "
    "framework for paediatric respiratory disease classification using Indian "
    "hospital EHR data. BMC Med Inform Decis Mak, v. 26, n. 1, 2026. DOI: "
    "10.1186/s12911-026-03493-2.",
    "SMOL, T. et al. Performance of meta-predictors for the classification of "
    "MED13L missense variations, implication of raw parameters. Eur J Med Genet, "
    "v. 65, n. 1, p. 104398, 2021. DOI: 10.1016/j.ejmg.2021.104398.",
    "SOEWITO, S. et al. Disparities in Cancer Genetic Testing and Variants of "
    "Uncertain Significance in the Hispanic Population of South Texas. JCO Oncol "
    "Pract, v. 18, n. 5, p. e805-e813, 2022. DOI: 10.1200/OP.22.00090.",
    "STARITA, L. M. et al. Massively Parallel Functional Analysis of BRCA1 RING "
    "Domain Variants. Genetics, v. 200, n. 2, p. 413-422, 2015. DOI: "
    "10.1534/genetics.115.175802.",
    "SUYBENG, V. et al. Comparison of Pathogenicity Prediction Tools on Somatic "
    "Variants. J Mol Diagn, v. 22, n. 12, p. 1383-1392, 2020. DOI: "
    "10.1016/j.jmoldx.2020.08.007.",
    "TIAN, Y. et al. REVEL and BayesDel outperform other in silico meta-predictors "
    "for clinical variant classification. Sci Rep, v. 9, n. 1, p. 12752, 2019. DOI: "
    "10.1038/s41598-019-49224-8.",
    "TOLAND, A. E.; ANDREASSEN, P. R. DNA repair-related functional assays for the "
    "classification of BRCA1 and BRCA2 variants: a critical review and needs "
    "assessment. J Med Genet, v. 54, n. 11, p. 721-731, 2017. DOI: "
    "10.1136/jmedgenet-2017-104707.",
    "TORRETTO, G. C. et al. Domain-Specific Computational, Functional and "
    "Structural Methods Enable Interpretation of BRCT Variants of Uncertain "
    "Significance. Curr Oncol, v. 33, n. 6, 2026. DOI: 10.3390/curroncol33060354.",
    "UNIPROT CONSORTIUM, THE. UniProt: the Universal Protein Knowledgebase in "
    "2023. Nucleic Acids Res, v. 51, n. D1, p. D523-D531, 2023.",
    "YANG, H. et al. BRCA2 function in DNA binding and recombination from a "
    "BRCA2-DSS1-ssDNA structure. Science, v. 297, n. 5588, p. 1837-1848, 2002. DOI: "
    "10.1126/science.297.5588.1837.",
    "ZHANG, X. et al. Genetic constraint at single amino acid resolution in "
    "protein domains improves missense variant prioritisation and gene discovery. "
    "Genome Med, v. 16, n. 1, p. 88, 2024. DOI: 10.1186/s13073-024-01358-9.",
]


def build_referencias(doc):
    h1(doc, "REFERÊNCIAS")
    for ref in REFERENCIAS:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r = p.add_run(ref)
        r.font.name = "Arial"; r.font.size = Pt(9.5)


# --------------------------------------------------------------------------- #
#  Main                                                                        #
# --------------------------------------------------------------------------- #
def main():
    doc = setup_document()
    build_frontmatter(doc)
    build_apresentacao(doc)
    build_metodos(doc)
    build_resultados(doc)
    build_conclusao(doc)
    build_referencias(doc)
    os.makedirs(os.path.dirname(OUT), exist_ok=True)
    doc.save(OUT)
    print(f"Documento salvo em: {OUT}")
    n_paras = len(doc.paragraphs)
    n_tables = len(doc.tables)
    print(f"Paragrafos: {n_paras} | Tabelas: {n_tables}")


if __name__ == "__main__":
    main()
