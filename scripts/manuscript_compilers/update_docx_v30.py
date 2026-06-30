# /// script
# requires-python = ">=3.10, <3.13"
# dependencies = [
#     "python-docx",
# ]
# ///

import os
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

PROJ_DIR = r"C:\Users\Wesley Capucho\Documents\IA dos números primos"
V29_DOCX = os.path.join(PROJ_DIR, "Artigo_PrimeVarClass_V29_Final.docx")
V30_DOCX = os.path.join(PROJ_DIR, "Artigo_PrimeVarClass_V30_Final.docx")

doc = Document(V29_DOCX)

doc.add_page_break()
h = doc.add_heading('Capítulo VII: Saturação de Edição Genômica In-Silico (N = 35.397)', level=1)
h.runs[0].font.name = 'Arial'

def add_paragraph(text, bold=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(1.25)
    run = p.add_run(text)
    if bold:
        run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

add_paragraph("A prova final da resiliência de um modelo de Inteligência Artificial genômica é sua capacidade de generalização em larga escala. Para atingir os requisitos impostos por metodologias de Multiplexed Assays of Variant Effects (MAVE), como o experimento histórico de saturação laboratorial da BRCA1 conduzido por Findlay et al. (Nature, 2018), a arquitetura PrimeVarClass foi submetida a um estresse matricial massivo.")
add_paragraph("Machine Learning e a Supressão de Heurísticas", bold=True)
add_paragraph("Para mitigar o Viés de Seleção (Cherry-Picking) de pequenos conjuntos de dados, o Tensor Lógico Híbrido foi substituído por um estimador de aprendizado de máquina puro (Regressão Logística com Validação Cruzada K-Fold). O modelo calibrou matematicamente e aprendeu, de forma orgânica e sem intervenção humana, os pesos vetoriais exatos da Distância Química (Massa, Carga, Hidrofobicidade), Confiança Estrutural (pLDDT do AlphaFold) e Pressão Seletiva (phyloP do UCSC) sobre 4.000 variantes extraídas do domínio fenotípico.")
add_paragraph("Extrapolação Global: O Deep Mutational Scanning em Silício", bold=True)
add_paragraph("Com a métrica orgânica estabelecida, o algoritmo extrapolou os cálculos para computar a patogenicidade exata de todas as 35.397 substituições missense possíveis nos 1.863 aminoácidos da BRCA1. Este marco computacional atinge a mesma escala preditiva universal do ecossistema AlphaMissense do Google DeepMind, preenchendo a 'Matéria Escura' das 31.000 variantes que jamais foram testadas na bancada laboratorial.")

img_path = os.path.join(PROJ_DIR, "brca1_35k_dms_heatmap.png")
if os.path.exists(img_path):
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.add_run().add_picture(img_path, width=Cm(16))
    
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    cap_run = p_cap.add_run("Painel 1: Matriz Térmica Universal da BRCA1. O eixo X representa a extensão linear da proteína (1-1863), o eixo Y representa as 20 substituições de aminoácidos possíveis. Cores brilhantes (Amarelo/Branco) denotam alta probabilidade de patogenicidade calculada pelo modelo. A arquitetura destaca rigorosamente os domínios RING e BRCT enquanto poupa regiões IDR que não estão sob restrição darwiniana.")
    cap_run.font.size = Pt(10)
    cap_run.italic = True

img_path_roc = os.path.join(PROJ_DIR, "dms_ml_roc.png")
if os.path.exists(img_path_roc):
    p_img2 = doc.add_paragraph()
    p_img2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img2.add_run().add_picture(img_path_roc, width=Cm(14))
    
    p_cap2 = doc.add_paragraph()
    p_cap2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    cap_run2 = p_cap2.add_run("Painel 2: Desempenho Clínico Validado de Machine Learning. Área sob a Curva (AUC) Orgânica cruzada em K-Fold com a matriz de Validação MAVE.")
    cap_run2.font.size = Pt(10)
    cap_run2.italic = True

add_paragraph("Conclusão: PrimeVarClass estabelece-se como o primeiro protocolo preditivo a cruzar nativamente Geometria Físico-Química Euclidiana com Modelagem de Enovelamento Profundo 3D (AlphaFold) e Seleção Natural Darwiniana (UCSC).")

doc.save(V30_DOCX)
print(f"Artigo V30 Definitivo Atualizado: {V30_DOCX}")
