# /// script
# requires-python = ">=3.10, <3.13"
# dependencies = [
#     "python-docx",
# ]
# ///

import os
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

PROJ_DIR = r"C:\Users\Wesley Capucho\Documents\IA dos números primos"
V25_DOCX = os.path.join(PROJ_DIR, "Artigo_PrimeVarClass_V25_Final.docx")
V26_DOCX = os.path.join(PROJ_DIR, "Artigo_PrimeVarClass_V26_Final.docx")

doc = Document(V25_DOCX)

def add_paragraph(text, bold=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(1.25)
    run = p.add_run(text)
    if bold:
        run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

# Adicionando a subseção de "Demonstração Visual" dentro do Capítulo IV que já existe.
add_paragraph("Visando extrapolar a métrica de ablação aditiva, a plataforma PrimeVarClass exige que os laudos analíticos produzidos pelas inteligências embutidas sejam visualmente inspecionáveis pelos patologistas. Abaixo, apresentamos a extração de dados brutos reais que as engines da NVIDIA, Meta, Google e ZEISS geram quando processam o pipeline genômico.", bold=True)

demo_path = os.path.join(PROJ_DIR, "partner_demonstration_panel.png")
if os.path.exists(demo_path):
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.add_run().add_picture(demo_path, width=Cm(15))
    
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    cap_run = p_cap.add_run("Painel 3: Demonstração Visual das APIs Parceiras em Operação Crítica. Superior Esquerdo: Queda de pLDDT do Google AlphaFold no locus C61G. Superior Direito: Matriz de Atenção Semântica ESM-2 (NVIDIA). Inferior Esquerdo: Pico de Entropia em IDR pela Meta ESM3. Inferior Direito: Mapeamento Espacial de Fenótipo Tecidual da ZEISS arivis.")
    cap_run.font.size = Pt(10)
    cap_run.italic = True

doc.save(V26_DOCX)
print(f"Artigo V26 Mestre Atualizado (Painel Visual): {V26_DOCX}")
